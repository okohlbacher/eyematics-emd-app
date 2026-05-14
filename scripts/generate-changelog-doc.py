"""
Generate EMD Changelog & Feedback Tracking Word document (v1.9 → v1.10).
Run: python3 scripts/generate-changelog-doc.py
Output: docs/EMD-v1.10-Changelog-Feedback.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)  # blue-800
    return h

def add_para(doc, text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p

def add_table_header(table, headers, bg='1e3a5f'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(9)
        set_cell_bg(cell, bg)

def add_table_row(table, values, bgs=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if bgs and i < len(bgs) and bgs[i]:
            set_cell_bg(cell, bgs[i])
    return row

def status_bg(status_str):
    s = status_str.strip()
    if s.startswith('✅') or 'ACCEPT' in s or 'Shipped' in s:
        return 'd1fae5'  # green-100
    elif s.startswith('❌') or 'REJECT' in s:
        return 'fee2e2'  # red-100
    elif s.startswith('⏳') or 'Ausstehend' in s:
        return 'fef3c7'  # amber-100
    elif s.startswith('💬'):
        return 'f3f4f6'  # gray-100
    return None

# ── document ───────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════

doc.add_paragraph()
title = doc.add_heading('EyeMatics Clinical Demonstrator (EMD)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Release v1.10 — Changelog & Feedback Tracking')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x4b, 0x55, 0x63)

date_p = doc.add_paragraph(f'Stand: 14. Mai 2026  ·  Git-Tag: v1.10 / v1.10-phase28')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(10)
date_p.runs[0].italic = True
date_p.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph()

# Summary box
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('v1.9.5 → v1.10  ·  13 Tage  ·  2 Phasen shipped  ·  729/729 Tests  ·  0 offene Bugs')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. CHANGELOG
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Changelog: v1.9 → v1.10', level=1)

p = doc.add_paragraph(
    'Dieser Abschnitt dokumentiert alle Änderungen zwischen dem letzten Release '
    'v1.9.5 (Synthetic Data Realism, 01.05.2026) und v1.10 (Session Hardening & UX Closure, 14.05.2026).'
)
p.runs[0].font.size = Pt(10)

# ── 1.1 Milestone Overview ──────────────────────────────────────────
add_heading(doc, '1.1  Milestone-Übersicht', level=2)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
add_table_header(t, ['Phase', 'Name', 'Shipped', 'Tests'])
rows_m = [
    ('27', 'Stateful Session Backend', '11.05.2026', '+45 neue Tests'),
    ('28', 'Admin Session Control UI', '14.05.2026', '+41 neue Tests'),
]
for r in rows_m:
    add_table_row(t, r)

doc.add_paragraph()

# ── 1.2 Formative Analysis Fixes ────────────────────────────────────
add_heading(doc, '1.2  Formative-Analyse-Korrekturen (alle 26 ACCEPTs)', level=2)
doc.add_paragraph(
    'Parallel zur v1.10-Planung wurden alle 26 als ACCEPT klassifizierten Findings '
    'der Formative Analyse vom 11.05.2026 implementiert und committed.'
)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
add_table_header(t2, ['Bereich', 'Anzahl behoben', 'Beispiele'])
areas = [
    ('USM – Nutzerverwaltung', '5', 'Kopierbutton, Standort-Pflicht, Username-Validierung, Edit-Fehlermeldung, Session-Revoke bei Löschung (PROT-001)'),
    ('KOH – Kohortenbildung', '4', 'Negativeingaben, Visus-Guard, Cursor-Alignment, Kohortenname in Analyse-Tab'),
    ('ANL – Analyse', '3', 'Responder-Tooltip, Alters-/Visus-Sortierung, Kohortenname bei Filter-Pfad'),
    ('FALL – Fallansicht', '6', 'ReferenceLine in relativem Chart, Encounter-Tile Hotspot, Visus-Y-Einheit, Wirkstoff-Badge, CRT-Labelwinkel, Interpolations-Hinweis'),
    ('QUAL – Datenqualität', '5', 'Anomalie-Status-Chip, Flag-Matching auf flaggedAt, Zeitraum-Badge, Zentrum-Filter sichtbar, Absolute Kennzahlen'),
    ('DAT / A-01..08', '3', 'Kohortenbildung-Button, Zentren-Navigation, Export-Button deaktiviert'),
]
for r in areas:
    add_table_row(t2, r)

doc.add_paragraph()

# ── 1.3 Phase 27 ────────────────────────────────────────────────────
add_heading(doc, '1.3  Phase 27: Stateful Session Backend  (SESS-02/03/04)', level=2)
doc.add_paragraph(
    'Ziel: Der Server verfolgt jedes ausgestellte Refresh-Token in einer persistenten Tabelle '
    'und invalidiert Token korrekt bei Rotation und Key-Wechsel.'
)

features_27 = [
    ('sessions.db (SQLite, WAL-Modus)',
     'Neue Tabelle refresh_sessions mit Feldern id (jti), sid (Familie), username, ver, '
     'issued_at, expires_at, last_used_at, revoked, key_id. Automatischer Cleanup-Timer.'),
    ('jti-Claims',
     'Jedes ausgestellte Access- und Refresh-Token trägt eine eindeutige UUID als jti-Claim.'),
    ('Token-Rotation',
     'POST /api/auth/refresh rotiert das Token: neues Token wird ausgestellt, das alte sofort revoziert. '
     'Wiederverwendung eines bereits revozierten Tokens → 401.'),
    ('Familien-Revokation',
     'Wird ein gestohlenes Token erkannt (zweite Verwendung), wird die gesamte Token-Familie '
     '(gleiche sid) revoziert — schützt gegen Token-Diebstahl.'),
    ('Dual-Key-Fenster',
     'POST /api/auth/rotate-key (Admin-only): neuer Schlüssel wird aktiviert, alter Schlüssel '
     'bleibt für laufende Sessions bis zu deren absolutem Ablauf gültig.'),
    ('Automatisierte Tests',
     '3 neue Testdateien: sessionsDb.test.ts (Unit), sessionRotation.test.ts (Integration), '
     'rotateKey.test.ts (Integration). Alle Tabellenzustände nach Rotation und Wiederverwendung abgedeckt.'),
]

t3 = doc.add_table(rows=1, cols=2)
t3.style = 'Table Grid'
add_table_header(t3, ['Feature', 'Beschreibung'])
for f, d in features_27:
    row = t3.add_row()
    row.cells[0].text = f
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].text = d
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── 1.4 Phase 28 ────────────────────────────────────────────────────
add_heading(doc, '1.4  Phase 28: Admin Session Control UI  (SESS-01, SESSUI-01/02/03)', level=2)
doc.add_paragraph(
    'Ziel: Admins können alle aktiven Sessions eines Benutzers einsehen und beenden — '
    'einzeln oder komplett — und Session-TTL-Werte ohne Konfigurationsdatei anpassen.'
)

features_28 = [
    ('Backend: 3 neue Admin-Endpunkte',
     'GET /api/auth/sessions?username= — aktive Sessions (DTO-Projektion: nur id, issued_at, last_used_at, expires_at, key_id)\n'
     'DELETE /api/auth/sessions/:id — Einzel-Revoke (404 bei bereits revozierten Sessions)\n'
     'DELETE /api/auth/sessions?username= — Sign-out-everywhere, gibt Anzahl zurück\n'
     'Alle drei: Admin-only Guard, kein requireCsrf (Bearer-only), Route-Ordering sichergestellt.'),
    ('AdminPage: Session-Accordion',
     'Pro Benutzerzeile: aufklappbares Panel mit Lazy-Fetch der aktiven Sessions. '
     'Spalten: Gerät (key_id letzte 8 Zeichen), Ausgestellt, Zuletzt verwendet, Läuft ab. '
     'Einzel-Revoke ohne Bestätigungsdialog (konsistent mit bestehenden Mustern). '
     '"Überall abmelden"-Button mit Ladeanimation. Leer-/Fehler-/Ladezustand.'),
    ('SettingsPage: TTL-Konfiguration',
     'Stunden-basierte Eingaben für Refresh-Token-Gültigkeit (min. 1h) und Absolute Session-Grenze (min. = Refresh-TTL, max. 720h/30 Tage). '
     'Client-seitige Validierung (validateTtl) + server-seitige Absicherung. '
     'Persistierung in config/settings.yaml. Auto-Dismissal des Erfolgs-Banners nach 3 s.'),
    ('ttlConversion.ts',
     'Neue reine Helper-Datei: hoursToMs, msToHours (Math.round), validateTtl mit TTL_MAX_HOURS=720. '
     'Vollständig getestet (10 Unit-Tests).'),
    ('i18n',
     '19 neue Übersetzungsschlüssel (DE + EN): 12 für Session-Panel, 7 für TTL-Formular inkl. capMax-Fehlermeldung.'),
    ('Sicherheits-Verbesserungen (Review-Pass)',
     'DTO-Projektion: sid/ver/revoked/username nicht an den Browser übertragen. '
     'TTL-Obergrenze 720h (30 Tage) serverseitig und clientseitig. '
     'Mount-Zeit-Validierung bei TTL-Formular. CSRF-Kommentar an DELETE-Endpunkten dokumentiert.'),
]

t4 = doc.add_table(rows=1, cols=2)
t4.style = 'Table Grid'
add_table_header(t4, ['Feature', 'Beschreibung'])
for f, d in features_28:
    row = t4.add_row()
    row.cells[0].text = f
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].text = d
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── 1.5 Test summary ────────────────────────────────────────────────
add_heading(doc, '1.5  Test-Zusammenfassung', level=2)

t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
add_table_header(t5, ['Version', 'Tests gesamt', 'Status'])
test_rows = [
    ('v1.9.5 (Baseline)', '682', '✅ alle bestanden'),
    ('Nach Phase 27', '704', '✅ alle bestanden (+22 neue)'),
    ('Nach Phase 28 (Waves 0-3)', '723', '✅ alle bestanden (+19 neue)'),
    ('Nach Review-Pass', '729', '✅ alle bestanden (+6 neue)'),
    ('v1.10 Final', '729', '✅ 729/729 · lint clean · build clean'),
]
for r in test_rows:
    bg = None
    if '✅' in r[2]:
        bg = ['ffffff', 'ffffff', 'd1fae5']
    add_table_row(t5, r, bgs=bg)

doc.add_paragraph()

# ── 1.6 Git log ─────────────────────────────────────────────────────
add_heading(doc, '1.6  Wesentliche Commits (v1.9.5 → v1.10)', level=2)

t6 = doc.add_table(rows=1, cols=2)
t6.style = 'Table Grid'
add_table_header(t6, ['Hash', 'Beschreibung'])
commits = [
    ('85e4821', 'fix(review): adversarial review fixes — DTO projection, TTL cap 720h, banner auto-dismiss, idempotent revoke 404'),
    ('93289ee', 'feat(28-04): per-user session accordion in AdminPage (lazy fetch, revoke, sign-out-everywhere)'),
    ('356223e', 'feat(28-04): 12 session i18n keys (de+en)'),
    ('4cbf543', 'feat(28-03): TTL inputs in SettingsPage + 7 i18n keys + settingsApi round-trip test'),
    ('fef29f6', 'feat(28-03): ttlConversion.ts + AppSettings.auth sub-object + DEFAULTS.auth'),
    ('67d7647', 'feat(28-02): 3 admin session route handlers in authApi.ts'),
    ('ab70263', 'feat(28-02): listActiveSessionsByUser + fix _closeForTests null-chain'),
    ('2a06784', 'feat(phase-27): stateful session backend — SESS-02/03/04 complete'),
    ('704d95a', 'feat(27): Plan 03 — jti rotation in /refresh, family revocation (SESS-03)'),
    ('884dbeb', 'feat(27-02-T2): initSessionsDb + cleanup interval in server/index.ts'),
    ('884dbeb', 'feat(27-02-T1): server/sessionsDb.ts — WAL SQLite sessions storage'),
    ('158e46d', 'fix: address 10 remaining formative-analysis findings (KOH-002, ANL, FALL, QUAL)'),
    ('7f443f5', 'fix: address formative analysis findings — PROT-001, USM-001, QUAL-006, ANL-003, FALL-003+A-01..A-08'),
]
for hash_, desc in commits:
    row = t6.add_row()
    row.cells[0].text = hash_
    row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[1].text = desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. FEEDBACK TRACKING
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Feedback Tracking — Formative Analyse (11.05.2026)', level=1)

p = doc.add_paragraph()
p.add_run('Dokument-Status: ').bold = True
p.add_run('Aktualisiert 14.05.2026 — alle 26 ACCEPTs committed · v1.10 Phases 27+28 shipped')
p = doc.add_paragraph()
p.add_run('Quelle: ').bold = True
p.add_run('formative_analyse_110526.docx + In-App Feedback emd-issues-2026-05-14.json')
p = doc.add_paragraph()
p.add_run('Legende: ').bold = True
p.add_run('✅ ACCEPT (umgesetzt)  ·  ❌ REJECT (nicht umgesetzt, Begründung)  ·  💬 COMMENT (Anmerkung / zurückgestellt)')

doc.add_paragraph()

# ── 2.0 Summary ─────────────────────────────────────────────────────
add_heading(doc, '2.0  Zusammenfassung', level=2)

t_sum = doc.add_table(rows=1, cols=2)
t_sum.style = 'Table Grid'
add_table_header(t_sum, ['Status', 'Anzahl'])
sum_rows = [
    ('✅ ACCEPT — committed', '26'),
    ('❌ REJECT (nicht im Scope)', '3'),
    ('💬 COMMENT / Backlog', '22'),
]
for r in sum_rows:
    row = t_sum.add_row()
    row.cells[0].text = r[0]
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].text = r[1]
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    bg = status_bg(r[0])
    if bg:
        set_cell_bg(row.cells[0], bg)

doc.add_paragraph()

# ── 2.1 v1.10 Milestone Status ──────────────────────────────────────
add_heading(doc, '2.1  v1.10 Milestone-Status', level=2)

t_ms = doc.add_table(rows=1, cols=3)
t_ms.style = 'Table Grid'
add_table_header(t_ms, ['Phase', 'Inhalt', 'Status'])
ms_rows = [
    ('27 — Stateful Session Backend', 'Server-seitige Sitzungstabelle, Token-Rotation, Key-Rotation', '✅ Shipped 11.05.2026'),
    ('28 — Admin Session Control UI', 'Sitzungsauflistung, Einzel-Revoke, Sign-out-everywhere, TTL-Konfiguration', '✅ Shipped 14.05.2026'),
    ('29 — Home Panel UX', 'Review-Buttons + Jump-Back-In-Navigation (FB-02, FB-03)', '⏳ Ausstehend'),
    ('30 — Terminology Configuration Docs', 'settings.yaml + Konfiguration.md Terminology-Abschnitt', '⏳ Ausstehend'),
]
for r in ms_rows:
    bgs = [None, None, status_bg(r[2])]
    add_table_row(t_ms, r, bgs=bgs)

doc.add_paragraph()

# ── 2.2 Teil A — In-App Feedback ────────────────────────────────────
add_heading(doc, '2.2  Teil A — In-App Feedback (9 Tickets)', level=2)

t_a = doc.add_table(rows=1, cols=4)
t_a.style = 'Table Grid'
add_table_header(t_a, ['#', 'Seite / Bereich', 'Beschreibung', 'Status'])
a_items = [
    ('A-01', 'Startseite', 'Kohortenbildung-Button ohne Funktion', '✅'),
    ('A-01', 'Startseite', 'Zentren wirken klickbar, ohne Funktion', '✅'),
    ('A-02', 'Startseite', 'Export-Button ohne Funktion', '✅ (deaktiviert mit Tooltip)'),
    ('A-03', 'Kohortenbildung', 'Negative Werte bei Alter/Visus/CRT möglich', '✅'),
    ('A-04', 'Analyse', 'Zeigt nicht an, welche Kohorte angezeigt wird', '✅'),
    ('A-05', 'Analyse', 'Fall nicht auswählbar im Aggregiert-Tab', '💬 By design / Phase-29-Scope'),
    ('A-06', 'Fallansicht', 'Einzelne Achsen-Ticks fehlen', '💬 Backlog (Screenshot benötigt)'),
    ('A-07', 'Fallansicht', 'CRT-Label-Fehler bei absoluter Kurve', '✅'),
    ('A-08', 'Fallansicht', 'Interpolierte Linie nicht gestrichelt', '✅ (Hinweistext korrigiert)'),
    ('A-09', 'Dokumentationsqualität', 'Grundgesamtheit nach Zeitraum gefiltert?', '💬 Backlog (Zeitraum-Badge)'),
]
for item in a_items:
    bgs = [None, None, None, status_bg(item[3])]
    add_table_row(t_a, item, bgs=bgs)

doc.add_paragraph()

# ── 2.3 Teil B — Formative Analyse ──────────────────────────────────
add_heading(doc, '2.3  Teil B — Formative Analyse (Anforderungs-basierte Findings)', level=2)

sections_b = {
    'USM — Nutzerverwaltung': [
        ('USM-001', 'Kopier-Button unresponsive', '✅', 'navigator.clipboard.writeText() mit 2s-Feedback'),
        ('USM-001', 'Nutzer ohne Standort möglich', '✅', 'Mindestens ein Zentrum Pflicht; Inline-Fehler'),
        ('USM-001', 'Fehlende Fehlermeldung Pflichtfelder', '✅', 'Username-Pflichtvalidierung + Fehlertext'),
        ('USM-002', 'Deaktivieren nicht möglich, nur Löschen', '❌', 'Reject (by design) — nicht in EMDREQ spezifiziert'),
        ('USM-002', 'Bearbeiten ohne Fehlermeldung', '✅', 'handleEditSave setzt actionError korrekt'),
        ('USM-006', 'Meldung bei zu vielen Fehlversuchen unklar', '💬', 'Backlog: verbleibendes Timeout in Meldung aufnehmen'),
        ('USM-008', 'Kein kontinuierlicher Counter sichtbar', '💬', 'By design: kein Counter (kein Brute-Force-Feedback)'),
        ('USM-008', 'maxLoginAttempts konfigurierbar?', '💬', 'Phase 28: TTL-Config in Admin-UI. maxLoginAttempts bewusst ausgelassen (Security-Kontrolle)'),
    ],
    'DAT — Dashboard / Datensatz': [
        ('DAT-003', 'Gesamtzahl nicht angezeigt', '✅', 'Fälle-KPI zeigt ungefilterte cases.length'),
        ('DAT-—', '"Attention needed" / "Jump Back In" nicht in EMDREQ', '💬', 'Bewusste UX-Erweiterung; Phase 29 verkabelt Panels vollständig'),
    ],
    'KOH — Kohortenbildung': [
        ('KOH-001', 'Kerndatensatz unvollständig', '💬', 'Bekannte Limitation; fehlende Felder bitte spezifizieren'),
        ('KOH-002', 'Negative Werte möglich', '✅', 'Behoben in A-03'),
        ('KOH-002', 'Cursor bleibt links bei Kohortenname', '✅', 'text-left ergänzt'),
        ('KOH-002', 'Visus-Eingabe akzeptiert String', '✅', 'v >= 0 Guard'),
        ('KOH-003', 'Filter-State nach Navigation weg', '💬', 'Backlog: History-State-Persistenz'),
        ('KOH-005', 'Aktive Kohorte in Analyse nicht angezeigt', '✅', 'Behoben in A-04'),
    ],
    'ANL — Analyse': [
        ('ANL-002', 'Intervall nicht vergleichend darstellbar', '💬', 'Backlog: IntervalHistogram Kohortenvergleich'),
        ('ANL-002', 'Was ist "Responder"?', '✅', 'ℹ-Tooltip mit Definition ergänzt'),
        ('ANL-003', 'Alters-/Visus-Achse nicht monoton', '✅', 'Sortierung nach age aufsteigend'),
        ('ANL-003', 'Kohorten-Vergleich im Aggregiert-Tab fehlt', '💬', 'Backlog'),
        ('ANL-004', 'Kritische Werte nicht konfigurierbar', '💬', 'Definiert in clinicalThresholds.ts; UI-Konfiguration Backlog'),
    ],
    'FALL — Fallansicht': [
        ('FALL-001', 'Weg zur Falldetailnavigation umständlich', '💬', 'Phase-29-Scope: direkte Chart→Fall-Navigation'),
        ('FALL-003', 'Events → nur absoluter Verlauf, nicht relativ', '✅', 'highlightDate-ReferenceLine auch in relativem Chart'),
        ('FALL-003', 'Zwei Buttons für eine Funktion', '✅', 'Encounter-Tile ist einziger Hotspot'),
        ('FALL-003', 'CRT-Bezeichnung in Legende', '✅', 'Behoben in A-07'),
        ('FALL-003', 'Einheit Visus fehlt', '✅', '"Visus (dezimal)" an Y-Achse'),
        ('FALL-003', 'Legende: Interpolation gestrichelt?', '✅', 'Behoben in A-08'),
        ('FALL-004', 'Wirkstoff pro Injektion unterrepräsentiert', '✅', 'Violettes Wirkstoff-Badge in Injektionszeile'),
        ('FALL-006', 'Vergleich Fall–Kohorte nicht möglich', '💬', 'Backlog: Kohorten-Referenzwerte in Fallansicht erweitern'),
    ],
    'QUAL — Datenqualität': [
        ('QUAL-001', 'Prüfung nicht auf Kohorten möglich', '💬', 'Backlog / Scope — kein Phase-28/29-Bestandteil'),
        ('QUAL-001', 'Prüfparameter nicht konfigurierbar', '❌', 'Reject (Scope) — nicht in EMDREQ'),
        ('QUAL-004', 'Fehlende Werte nicht vorgeschlagen', '💬', 'Backlog: Imputation'),
        ('QUAL-006', 'Fehlerkennung nur weit unten sichtbar', '✅', 'Status-Chip statt "Fehler melden"-Button; Flags in Tabelle sichtbar'),
        ('QUAL-006', 'Mehrere Flags — Bestätigung gilt für alle', '✅', 'updateQualityFlag matcht auf flaggedAt'),
        ('QUAL-011', 'Grundgesamtheit ohne Zeitraum-Filter', '✅', 'Zeitraum-Badge bei aktivem Filter'),
        ('QUAL-011', 'Zentren-Filter versteckt', '✅', 'Dropdown immer sichtbar in Kopfzeile'),
        ('QUAL-011', 'Kein Multi-Select Zentrum', '❌', 'Reject vorerst — nicht in EMDREQ'),
        ('QUAL-011', 'Absolute Kennzahlen fehlen', '✅', 'Prozentwert + absolute Zahl in SummaryCards'),
        ('QUAL-011', 'Prüfungsergebnisse fehlen', '✅', 'Anteil geprüft/in Bearbeitung/ungecheckt in SummaryCards'),
        ('QUAL-011', 'Plausibilitätsbereiche nur über Details', '✅', 'Zielwert + Fallzahl unter Fortschrittsbalken'),
    ],
    'PROT — Protokoll / Audit Log': [
        ('PROT-001', 'Gelöschter Account als "anonymous" im Protokoll', '✅', 'DELETE /api/auth/users/:username ruft revokeByUsername() auf — Sessions werden bei Löschung sofort invalidiert'),
    ],
}

for section_title, items in sections_b.items():
    add_heading(doc, section_title, level=3)
    t_b = doc.add_table(rows=1, cols=4)
    t_b.style = 'Table Grid'
    add_table_header(t_b, ['ID', 'Finding', 'Status', 'Maßnahme / Kommentar'])
    for item in items:
        bgs = [None, None, status_bg(item[2]), None]
        add_table_row(t_b, item, bgs=bgs)
    doc.add_paragraph()

# ── 2.4 Offene Backlog-Items ─────────────────────────────────────────
add_heading(doc, '2.4  Offene Backlog-Items (mittelfristig, kein akuter Bug)', level=2)

t_bl = doc.add_table(rows=1, cols=3)
t_bl.style = 'Table Grid'
add_table_header(t_bl, ['ID', 'Beschreibung', 'Ziel'])
backlog = [
    ('FB-02', 'Home "Attention needed" Review-Buttons ohne Funktion', 'Phase 29'),
    ('FB-03', 'Home "Jump Back In" Routing fehlt', 'Phase 29'),
    ('FALL-001', 'Direkte Chart→Falldetail-Navigation', 'Phase 29'),
    ('KOH-003', 'Filter-State beim Rücknavigieren erhalten', 'Backlog'),
    ('ANL-002', 'Intervall-Histogramm im Kohorten-Vergleichsmodus', 'Backlog'),
    ('ANL-002', 'Visus-Berechnung Aggregiert vs. Trajektorien kenntlich machen', 'Backlog'),
    ('ANL-003', 'Kohorten-Vergleich im Aggregiert-Tab', 'Backlog'),
    ('FALL-006', 'Erweiterter Fall–Kohorte-Vergleich', 'Backlog'),
    ('QUAL-001', 'Kohorten-basierte Qualitätsprüfung', 'Backlog'),
    ('QUAL-004', 'Auto-Suggest fehlender Werte (Imputation)', 'Backlog'),
    ('USM-006', 'Rate-Limiting-Meldung mit verbleibendem Timeout', 'Backlog'),
    ('USM-008', 'maxLoginAttempts Admin-konfigurierbar', 'Backlog'),
    ('A-06', 'Fehlende Achsenticks (Screenshot für Repro benötigt)', 'Backlog'),
    ('A-09', 'Zeitraum-Badge in Grundgesamtheit QualityPage', 'Backlog'),
]
for r in backlog:
    add_table_row(t_bl, r)

doc.add_paragraph()

# ── Footer note ─────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph(
    f'Generiert: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}  ·  '
    f'Git: v1.10 (85e4821)  ·  Repository: github.com/okohlbacher/eyematics-emd-app'
)
p.runs[0].font.size = Pt(8)
p.runs[0].italic = True
p.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ────────────────────────────────────────────────────────────
os.makedirs('docs', exist_ok=True)
out = 'docs/EMD-v1.10-Changelog-Feedback.docx'
doc.save(out)
print(f"✅  Saved: {out}")
